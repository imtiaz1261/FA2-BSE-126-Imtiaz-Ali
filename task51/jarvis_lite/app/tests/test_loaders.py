"""Tests for the document loaders and the loader factory."""

import docx
import pytest

from app.core.exceptions import DocumentLoadError, EmptyDocumentError, UnsupportedFileTypeError
from app.loaders.loader_factory import get_loader, load_document
from app.loaders.txt_loader import TXTLoader
from app.loaders.docx_loader import DOCXLoader


def test_txt_loader_reads_content(tmp_path):
    file_path = tmp_path / "notes.txt"
    file_path.write_text("Hello world.\n\nSecond paragraph.", encoding="utf-8")

    documents = TXTLoader().load(str(file_path))

    assert len(documents) == 1
    assert "Hello world." in documents[0].content
    assert documents[0].metadata["filename"] == "notes.txt"
    assert documents[0].metadata["file_type"] == "txt"


def test_txt_loader_empty_file_returns_no_documents(tmp_path):
    file_path = tmp_path / "empty.txt"
    file_path.write_text("   \n\n  ", encoding="utf-8")

    assert TXTLoader().load(str(file_path)) == []


def test_docx_loader_reads_paragraphs(tmp_path):
    file_path = tmp_path / "report.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(str(file_path))

    documents = DOCXLoader().load(str(file_path))

    assert len(documents) == 1
    assert "First paragraph." in documents[0].content
    assert "Second paragraph." in documents[0].content
    assert documents[0].metadata["file_type"] == "docx"


def test_loader_factory_picks_correct_loader(tmp_path):
    file_path = tmp_path / "file.txt"
    file_path.write_text("content", encoding="utf-8")

    loader = get_loader(str(file_path))

    assert isinstance(loader, TXTLoader)


def test_loader_factory_rejects_unsupported_extension(tmp_path):
    file_path = tmp_path / "audio.mp3"
    file_path.write_bytes(b"not really audio")

    with pytest.raises(UnsupportedFileTypeError):
        get_loader(str(file_path))


def test_load_document_missing_file_raises():
    with pytest.raises(DocumentLoadError):
        load_document("/nonexistent/path/file.txt")


def test_load_document_empty_file_raises(tmp_path):
    file_path = tmp_path / "blank.txt"
    file_path.write_text("   ", encoding="utf-8")

    with pytest.raises(EmptyDocumentError):
        load_document(str(file_path))
