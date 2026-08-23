"""DOCX loader — joins non-empty paragraphs into a single LoadedDocument."""

import logging
from pathlib import Path
from typing import List

import docx
from docx.opc.exceptions import PackageNotFoundError

from app.core.exceptions import DocumentLoadError
from app.loaders.base import BaseLoader, LoadedDocument

logger = logging.getLogger(__name__)


class DOCXLoader(BaseLoader):
    def load(self, file_path: str) -> List[LoadedDocument]:
        path = Path(file_path)
        try:
            document = docx.Document(str(path))
        except (PackageNotFoundError, OSError) as exc:
            logger.exception("Failed to open DOCX %s", file_path)
            raise DocumentLoadError(f"Could not open DOCX '{path.name}': {exc}") from exc

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        if not content.strip():
            logger.info("DOCX %s contained no extractable text", path.name)
            return []

        return [
            LoadedDocument(
                content=content,
                metadata={"source": str(path), "filename": path.name, "file_type": "docx"},
            )
        ]
