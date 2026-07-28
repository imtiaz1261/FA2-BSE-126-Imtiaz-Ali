"""
report/report_builder.py
--------------------------
Builds the final structured research report (Markdown), and exports it
to PDF and DOCX.

Report sections (per spec):
  Title, Executive Summary, Research Objectives, Key Findings,
  Detailed Analysis, Supporting Evidence, References & Sources,
  Conclusion, Future Recommendations

The narrative sections (Executive Summary, Key Findings, Detailed
Analysis, Conclusion, Future Recommendations) are drafted by the LLM
in a single call, grounded strictly in the collected task summaries --
References & Sources is assembled programmatically from the actual
collected URLs (never LLM-generated) so citations can't be hallucinated.
"""

import re
from pathlib import Path
from typing import Tuple

from llm import get_llm
from state import ResearchState
from utils import get_logger

logger = get_logger(__name__)


def _collect_all_sources(state: ResearchState):
    """Flatten and de-duplicate (by URL) every source collected across tasks."""
    seen = set()
    sources = []
    for task, results in state.get("search_results", {}).items():
        for r in results:
            url = r.get("url", "")
            if url and url in seen:
                continue
            if url:
                seen.add(url)
            sources.append(r)
    return sources


def _draft_narrative_sections(state: ResearchState) -> str:
    """Single LLM call producing the five narrative sections as Markdown."""
    combined_summary = state.get("combined_summary", "")
    objective = state.get("objective", state.get("query", ""))

    prompt = (
        "You are writing a research report based ONLY on the verified summary "
        "below -- do not add outside knowledge or invent facts/statistics not "
        "present in the summary.\n\n"
        f"Research objective: {objective}\n\n"
        f"Verified findings:\n{combined_summary}\n\n"
        "Write the following sections as Markdown, using '## <Section Name>' "
        "headers exactly as given, in this order:\n"
        "## Executive Summary\n"
        "## Key Findings\n"
        "## Detailed Analysis\n"
        "## Conclusion\n"
        "## Future Recommendations\n\n"
        "Key Findings should be a bullet list. Keep the whole response focused "
        "and factual."
    )
    response = get_llm().invoke(prompt)
    return response.content.strip()


def build_markdown_report(state: ResearchState) -> Tuple[str, str]:
    """Return (title, full_markdown_report)."""
    query = state.get("query", "Untitled Research")
    title = f"Research Report: {query}"

    narrative = _draft_narrative_sections(state)

    all_sources = _collect_all_sources(state)
    references_lines = [
        f"{i}. [{s['title']}]({s['url']})" if s.get("url") else f"{i}. {s['title']}"
        for i, s in enumerate(all_sources, start=1)
    ]
    references_md = "\n".join(references_lines) if references_lines else "_No sources were collected._"

    supporting_evidence_lines = []
    for task, summary in state.get("task_summaries", {}).items():
        supporting_evidence_lines.append(f"**{task}**\n\n{summary}\n")
    supporting_evidence_md = "\n".join(supporting_evidence_lines) or "_No supporting evidence available._"

    objectives_lines = "\n".join(f"- {t}" for t in state.get("tasks", []))

    errors = state.get("errors", [])
    errors_note = ""
    if errors:
        errors_note = "\n\n## Notes\n" + "\n".join(f"- {e}" for e in errors)

    full_markdown = (
        f"# {title}\n\n"
        f"{narrative}\n\n"
        f"## Research Objectives\n{objectives_lines}\n\n"
        f"## Supporting Evidence\n{supporting_evidence_md}\n\n"
        f"## References & Sources\n{references_md}"
        f"{errors_note}\n"
    )
    return title, full_markdown


# --------------------------------------------------------------------------
# Export: PDF (reportlab)
# --------------------------------------------------------------------------
def export_pdf(title: str, markdown_text: str, output_path: Path) -> Path:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.lib.styles import getSampleStyleSheet

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 16)]

    for raw_line in markdown_text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            story.append(Spacer(1, 8))
            continue
        if line.startswith("# "):
            continue  # title already added
        if line.startswith("## "):
            story.append(Spacer(1, 10))
            story.append(Paragraph(line[3:], styles["Heading2"]))
        elif line.strip().startswith(("- ", "* ")):
            text = re.sub(r"^[\-\*]\s+", "", line.strip())
            text = _markdown_inline_to_reportlab(text)
            story.append(ListFlowable(
                [ListItem(Paragraph(text, styles["Normal"]))], bulletType="bullet",
            ))
        elif re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s+", "", line.strip())
            text = _markdown_inline_to_reportlab(text)
            story.append(Paragraph(f"\u2022 {text}", styles["Normal"]))
        else:
            story.append(Paragraph(_markdown_inline_to_reportlab(line), styles["Normal"]))

    doc.build(story)
    logger.info("PDF report exported to %s", output_path)
    return output_path


def _markdown_inline_to_reportlab(text: str) -> str:
    """Convert minimal inline markdown (bold, links) to ReportLab-safe markup."""
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r'<link href="\2">\1</link>', text)
    return text


# --------------------------------------------------------------------------
# Export: DOCX (python-docx)
# --------------------------------------------------------------------------
def export_docx(title: str, markdown_text: str, output_path: Path) -> Path:
    import docx

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = docx.Document()
    document.add_heading(title, level=0)

    for raw_line in markdown_text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.strip().startswith(("- ", "* ")):
            text = re.sub(r"^[\-\*]\s+", "", line.strip())
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            document.add_paragraph(text, style="List Bullet")
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            document.add_paragraph(text)

    document.save(str(output_path))
    logger.info("DOCX report exported to %s", output_path)
    return output_path
