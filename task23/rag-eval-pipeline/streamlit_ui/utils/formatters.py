"""
utils/formatters.py
-------------------
Pure-function helpers for formatting, exporting, and rendering content.
No Streamlit imports — these are data-layer utilities only.
"""

from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any


# ──────────────────────────────────────────────────────────────────────────────
# Timestamp formatting
# ──────────────────────────────────────────────────────────────────────────────

def format_timestamp(dt: datetime, short: bool = False) -> str:
    """
    Human-readable timestamp.

    Parameters
    ----------
    dt    : datetime to format
    short : if True → "14:32", else → "Today at 14:32" / "Jul 27 at 14:32"
    """
    now = datetime.now()
    time_str = dt.strftime("%H:%M")

    if short:
        return time_str

    if dt.date() == now.date():
        return f"Today at {time_str}"
    elif (now.date() - dt.date()).days == 1:
        return f"Yesterday at {time_str}"
    else:
        return dt.strftime("%b %d at %H:%M")


def format_relative_time(dt: datetime) -> str:
    """
    Return a concise relative label, e.g. "2m ago", "3h ago", "Jul 25".
    """
    delta = datetime.now() - dt
    seconds = int(delta.total_seconds())

    if seconds < 60:
        return "just now"
    elif seconds < 3600:
        return f"{seconds // 60}m ago"
    elif seconds < 86400:
        return f"{seconds // 3600}h ago"
    elif seconds < 604800:
        return f"{delta.days}d ago"
    else:
        return dt.strftime("%b %d")


def format_duration(seconds: float) -> str:
    """Convert elapsed seconds to a readable string like '1m 23s' or '45s'."""
    if seconds < 60:
        return f"{seconds:.1f}s"
    mins = int(seconds) // 60
    secs = int(seconds) % 60
    return f"{mins}m {secs}s"


# ──────────────────────────────────────────────────────────────────────────────
# Export helpers
# ──────────────────────────────────────────────────────────────────────────────

def messages_to_markdown(messages: list[dict[str, Any]], title: str = "Chat Export") -> str:
    """
    Serialise a conversation message list to a Markdown string.

    Parameters
    ----------
    messages : list of message dicts with keys 'role', 'content', 'timestamp'
    title    : document heading
    """
    lines: list[str] = [
        f"# {title}",
        f"*Exported on {datetime.now().strftime('%B %d, %Y at %H:%M')}*",
        "",
        "---",
        "",
    ]

    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        ts   = msg.get("timestamp")
        ts_str = format_timestamp(ts) if isinstance(ts, datetime) else ""

        lines.append(f"### {role}  _{ts_str}_")
        lines.append("")
        lines.append(msg.get("content", ""))
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def messages_to_txt(messages: list[dict[str, Any]]) -> str:
    """Plain-text export — no markdown syntax."""
    lines: list[str] = []
    for msg in messages:
        role = msg.get("role", "?").upper()
        ts   = msg.get("timestamp")
        ts_str = f"[{format_timestamp(ts)}]" if isinstance(ts, datetime) else ""
        lines.append(f"{role} {ts_str}")
        lines.append(msg.get("content", ""))
        lines.append("")
    return "\n".join(lines)


def messages_to_json(messages: list[dict[str, Any]], title: str = "Chat Export") -> str:
    """JSON export — datetimes serialised to ISO strings."""
    payload = {
        "title":     title,
        "exported":  datetime.now().isoformat(),
        "messages":  [],
    }
    for msg in messages:
        item = {k: v for k, v in msg.items()}
        # Convert datetime → string
        if isinstance(item.get("timestamp"), datetime):
            item["timestamp"] = item["timestamp"].isoformat()
        payload["messages"].append(item)
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_to_bytes_txt(messages: list[dict[str, Any]]) -> bytes:
    return messages_to_txt(messages).encode("utf-8")


def export_to_bytes_markdown(messages: list[dict[str, Any]], title: str) -> bytes:
    return messages_to_markdown(messages, title).encode("utf-8")


def export_to_bytes_json(messages: list[dict[str, Any]], title: str) -> bytes:
    return messages_to_json(messages, title).encode("utf-8")


def export_to_bytes_docx(messages: list[dict[str, Any]], title: str) -> bytes | None:
    """
    Generate a .docx file. Returns None if python-docx is not installed
    (caller should warn the user and fall back to TXT).
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ImportError:
        return None

    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Exported on {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    doc.add_paragraph("")

    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        ts   = msg.get("timestamp")
        ts_str = format_timestamp(ts) if isinstance(ts, datetime) else ""

        doc.add_heading(f"{role}  —  {ts_str}", level=2)
        doc.add_paragraph(msg.get("content", ""))
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_to_bytes_pdf(messages: list[dict[str, Any]], title: str) -> bytes | None:
    """
    Generate a PDF using reportlab. Returns None if not installed.
    """
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
    except ImportError:
        return None

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2*cm, rightMargin=2*cm,
                             topMargin=2*cm,  bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))

    for msg in messages:
        role = msg.get("role", "?").capitalize()
        ts   = msg.get("timestamp")
        ts_str = format_timestamp(ts) if isinstance(ts, datetime) else ""
        story.append(Paragraph(f"<b>{role}</b>  <font size='9' color='grey'>{ts_str}</font>",
                                styles["Heading3"]))
        # Escape HTML special chars in content
        content = (msg.get("content", "")
                   .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        story.append(Paragraph(content, styles["BodyText"]))
        story.append(Spacer(1, 10))

    doc.build(story)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# Token / stat estimates (placeholder — replace with real tokeniser)
# ──────────────────────────────────────────────────────────────────────────────

def estimate_tokens(text: str) -> int:
    """Rough heuristic: ~4 characters per token (GPT-style)."""
    return max(1, len(text) // 4)


def conversation_stats(messages: list[dict[str, Any]]) -> dict[str, int]:
    """Return a dict of basic statistics about a conversation."""
    user_msgs      = [m for m in messages if m.get("role") == "user"]
    assistant_msgs = [m for m in messages if m.get("role") == "assistant"]
    total_tokens   = sum(estimate_tokens(m.get("content", "")) for m in messages)
    return {
        "total_messages":     len(messages),
        "user_messages":      len(user_msgs),
        "assistant_messages": len(assistant_msgs),
        "estimated_tokens":   total_tokens,
    }


# ──────────────────────────────────────────────────────────────────────────────
# Misc text utils
# ──────────────────────────────────────────────────────────────────────────────

def truncate(text: str, max_len: int = 60) -> str:
    """Truncate *text* to *max_len* characters, adding '…' if needed."""
    return text if len(text) <= max_len else text[:max_len].rstrip() + "…"


def sanitise_filename(name: str) -> str:
    """Strip characters that are unsafe in file names."""
    safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in name)
    return safe.strip("_").strip() or "export"
