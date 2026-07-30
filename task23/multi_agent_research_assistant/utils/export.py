"""
utils/export.py
---------------
Export helpers for chat conversations.

Supported formats
-----------------
- Markdown  (.md)   — plain text, always available
- TXT       (.txt)  — stripped, always available
- PDF       (.pdf)  — requires reportlab
- DOCX      (.docx) — requires python-docx

Each function returns ``bytes`` so Streamlit's ``st.download_button``
can consume them directly without writing to disk.

Usage (inside a Streamlit component)::

    from utils.export import export_markdown

    st.download_button(
        label="Download .md",
        data=export_markdown(messages, title),
        file_name="chat.md",
        mime="text/markdown",
    )
"""

from __future__ import annotations

import io
import re
import textwrap
from datetime import datetime, timezone
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from utils.session import Message


# ── Internal helpers ──────────────────────────────────────────────────────────

def _header(title: str) -> str:
    """Return a consistent plain-text header block."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    return f"# {title}\n\nExported: {now}\n\n---\n\n"


def _strip_markdown(text: str) -> str:
    """Very light Markdown → plain-text conversion (no external deps)."""
    text = re.sub(r"```[\s\S]*?```", lambda m: m.group(0).replace("```", ""), text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",   r"\1", text)
    text = re.sub(r"#+\s",        "",    text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def _messages_to_md(messages: list["Message"], title: str) -> str:
    """Convert a list of messages to a Markdown string."""
    lines: list[str] = [_header(title)]
    for msg in messages:
        role_label = "**You**" if msg["role"] == "user" else "**Assistant**"
        ts = msg.get("timestamp", "")
        lines.append(f"### {role_label}  `{ts}`\n")
        lines.append(msg["content"])
        lines.append("\n---\n")
    return "\n".join(lines)


# ── Public export functions ───────────────────────────────────────────────────

def export_markdown(messages: list["Message"], title: str = "Chat Export") -> bytes:
    """
    Serialise messages to a UTF-8 Markdown document.

    Parameters
    ----------
    messages : list[Message]
    title    : str  Conversation title used in the document heading.

    Returns
    -------
    bytes  UTF-8 encoded Markdown.
    """
    md = _messages_to_md(messages, title)
    return md.encode("utf-8")


def export_txt(messages: list["Message"], title: str = "Chat Export") -> bytes:
    """
    Serialise messages to a plain-text document (Markdown stripped).

    Returns
    -------
    bytes  UTF-8 encoded plain text.
    """
    md = _messages_to_md(messages, title)
    plain = _strip_markdown(md)
    return plain.encode("utf-8")


def export_pdf(messages: list["Message"], title: str = "Chat Export") -> bytes:
    """
    Serialise messages to a PDF document using reportlab.

    Falls back to a minimal error PDF if reportlab is unavailable.

    Returns
    -------
    bytes  PDF binary data.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        )
    except ImportError:
        return b"%PDF-1.4 % reportlab not installed\n"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.grey,
        spaceAfter=4,
    )
    user_style = ParagraphStyle(
        "User",
        parent=styles["Normal"],
        fontSize=11,
        backColor=colors.HexColor("#DCF8C6"),
        borderPadding=(6, 8, 6, 8),
        spaceAfter=6,
        leading=16,
    )
    ai_style = ParagraphStyle(
        "AI",
        parent=styles["Normal"],
        fontSize=11,
        backColor=colors.HexColor("#F0F4FF"),
        borderPadding=(6, 8, 6, 8),
        spaceAfter=6,
        leading=16,
    )

    story = [
        Paragraph(title, title_style),
        Paragraph(
            f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            meta_style,
        ),
        HRFlowable(width="100%", spaceAfter=12),
    ]

    for msg in messages:
        role_label = "You" if msg["role"] == "user" else "Assistant"
        ts = msg.get("timestamp", "")
        story.append(Paragraph(f"<b>{role_label}</b> &nbsp; <font size='8' color='grey'>{ts}</font>", meta_style))

        style = user_style if msg["role"] == "user" else ai_style
        # Escape HTML special chars for reportlab
        safe_content = (
            msg["content"]
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )
        story.append(Paragraph(safe_content, style))
        story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()


def export_docx(messages: list["Message"], title: str = "Chat Export") -> bytes:
    """
    Serialise messages to a DOCX document using python-docx.

    Falls back to UTF-8 plain text if python-docx is unavailable.

    Returns
    -------
    bytes  DOCX binary data (or plain text as fallback).
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return export_txt(messages, title)

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Export timestamp
    meta = doc.add_paragraph(
        f"Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    meta.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    meta.runs[0].font.size = Pt(9)

    doc.add_paragraph("─" * 60)

    for msg in messages:
        role_label = "You" if msg["role"] == "user" else "Assistant"
        ts = msg.get("timestamp", "")

        # Role + timestamp header
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f"{role_label}  ")
        role_run.bold = True
        role_run.font.size = Pt(11)
        ts_run = role_para.add_run(ts)
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        # Message body
        body_para = doc.add_paragraph(msg["content"])
        body_para.runs[0].font.size = Pt(11)
        if msg["role"] == "user":
            body_para.runs[0].font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)
        else:
            body_para.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        doc.add_paragraph()  # blank spacer

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Convenience map ───────────────────────────────────────────────────────────

EXPORT_OPTIONS: dict[str, dict] = {
    "Markdown (.md)": {
        "fn":        export_markdown,
        "file_name": "chat_export.md",
        "mime":      "text/markdown",
    },
    "Plain Text (.txt)": {
        "fn":        export_txt,
        "file_name": "chat_export.txt",
        "mime":      "text/plain",
    },
    "PDF (.pdf)": {
        "fn":        export_pdf,
        "file_name": "chat_export.pdf",
        "mime":      "application/pdf",
    },
    "Word (.docx)": {
        "fn":        export_docx,
        "file_name": "chat_export.docx",
        "mime":      "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    },
}
