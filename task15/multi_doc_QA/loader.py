"""
loader.py
---------
Document Loading + Metadata Extraction.

Responsible for reading raw files from the `data/` directory (or any
directory the caller specifies) and turning them into a list of
LangChain `Document` objects, each carrying rich metadata:

    - source        : full file path
    - file_name     : just the file name
    - file_type     : "pdf" | "docx" | "txt"
    - page          : page number (PDFs only, 1-indexed)
    - paragraph     : paragraph index (DOCX only, 1-indexed)

Supported formats: .pdf, .docx, .txt
Unsupported formats and unreadable/empty files are skipped with a
logged warning rather than crashing the whole pipeline.
"""

from pathlib import Path
from typing import List

from langchain_core.documents import Document
from pypdf import PdfReader
import docx  # python-docx

from config import SUPPORTED_EXTENSIONS
from utils import get_logger, is_non_empty_text

logger = get_logger(__name__)


class DocumentLoadError(Exception):
    """Raised when no valid documents could be loaded at all."""


# --------------------------------------------------------------------------
# Individual format loaders
# --------------------------------------------------------------------------
def _load_pdf(file_path: Path) -> List[Document]:
    """Load a PDF, producing one Document per non-empty page."""
    docs: List[Document] = []
    try:
        reader = PdfReader(str(file_path))
    except Exception as exc:
        logger.error("Failed to open PDF '%s': %s", file_path.name, exc)
        return docs

    if len(reader.pages) == 0:
        logger.warning("PDF '%s' has no pages, skipping.", file_path.name)
        return docs

    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            logger.warning(
                "Could not extract text from page %d of '%s': %s",
                page_number, file_path.name, exc,
            )
            continue

        if not is_non_empty_text(text):
            continue

        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "pdf",
                    "page": page_number,
                },
            )
        )

    if not docs:
        logger.warning("PDF '%s' contained no extractable text.", file_path.name)
    return docs


def _load_docx(file_path: Path) -> List[Document]:
    """Load a DOCX, producing one Document per non-empty paragraph group.

    Consecutive paragraphs are kept separate (with paragraph index in
    metadata) so downstream chunking can still recombine them via the
    text splitter while retaining traceability to the original section.
    """
    docs: List[Document] = []
    try:
        document = docx.Document(str(file_path))
    except Exception as exc:
        logger.error("Failed to open DOCX '%s': %s", file_path.name, exc)
        return docs

    for para_index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text
        if not is_non_empty_text(text):
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "docx",
                    "paragraph": para_index,
                },
            )
        )

    if not docs:
        logger.warning("DOCX '%s' contained no extractable text.", file_path.name)
    return docs


def _load_txt(file_path: Path) -> List[Document]:
    """Load a plain text file as a single Document."""
    try:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        logger.error("Failed to read TXT '%s': %s", file_path.name, exc)
        return []

    if not is_non_empty_text(text):
        logger.warning("TXT file '%s' is empty, skipping.", file_path.name)
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "txt",
            },
        )
    ]


_LOADER_MAP = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_txt,
}


# --------------------------------------------------------------------------
# Public API
# --------------------------------------------------------------------------
def load_document(file_path: Path) -> List[Document]:
    """Load a single file (any supported extension) into Documents."""
    extension = file_path.suffix.lower()
    loader_fn = _LOADER_MAP.get(extension)

    if loader_fn is None:
        logger.warning(
            "Unsupported file format '%s' for file '%s'. Skipping.",
            extension, file_path.name,
        )
        return []

    logger.info("Loading %s ...", file_path.name)
    docs = loader_fn(file_path)
    logger.info("Loaded %d section(s) from %s", len(docs), file_path.name)
    return docs


def load_documents_from_directory(directory: Path) -> List[Document]:
    """
    Load every supported file found in `directory` (non-recursive).

    Raises
    ------
    DocumentLoadError
        If the directory does not exist, is empty, or contains no
        documents that could be successfully parsed.
    """
    directory = Path(directory)

    if not directory.exists():
        raise DocumentLoadError(f"Directory not found: {directory}")

    files = sorted(p for p in directory.iterdir() if p.is_file())
    if not files:
        raise DocumentLoadError(f"No files found in directory: {directory}")

    supported_files = [f for f in files if f.suffix.lower() in SUPPORTED_EXTENSIONS]
    if not supported_files:
        raise DocumentLoadError(
            f"No supported files (.pdf, .docx, .txt) found in: {directory}"
        )

    all_docs: List[Document] = []
    for file_path in supported_files:
        all_docs.extend(load_document(file_path))

    if not all_docs:
        raise DocumentLoadError(
            "Files were found, but none contained readable text "
            "(they may be empty, corrupted, or scanned images)."
        )

    logger.info(
        "Finished loading directory '%s': %d document section(s) from %d file(s).",
        directory, len(all_docs), len(supported_files),
    )
    return all_docs