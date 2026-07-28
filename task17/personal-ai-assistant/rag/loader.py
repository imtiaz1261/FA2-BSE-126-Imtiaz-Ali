"""
rag/loader.py
-------------
Loads a single PDF, DOCX, or TXT file into LangChain Documents, with
metadata (file name, type, page/paragraph) preserved for traceability.
Used by tools/file_tool.py to support "summarize this PDF" / "what
does section 2 of my notes say" style requests.
"""

from pathlib import Path
from typing import List

from langchain_core.documents import Document
from pypdf import PdfReader
import docx

from utils import get_logger

logger = get_logger(__name__)


class FileLoadError(Exception):
    """Raised when a file cannot be found or parsed."""


def _load_pdf(file_path: Path) -> List[Document]:
    docs: List[Document] = []
    reader = PdfReader(str(file_path))
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "pdf",
                    "page": page_number,
                },
            )
        )
    return docs


def _load_docx(file_path: Path) -> List[Document]:
    docs: List[Document] = []
    document = docx.Document(str(file_path))
    for para_index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "docx",
                    "paragraph": para_index,
                },
            )
        )
    return docs


def _load_txt(file_path: Path) -> List[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "txt",
            },
        )
    ]


_LOADERS = {".pdf": _load_pdf, ".docx": _load_docx, ".txt": _load_txt}


def load_file(file_path: Path) -> List[Document]:
    """Load one supported file into a list of Documents."""
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileLoadError(f"File not found: {file_path}")

    extension = file_path.suffix.lower()
    loader_fn = _LOADERS.get(extension)
    if loader_fn is None:
        raise FileLoadError(
            f"Unsupported file type '{extension}'. Supported: .pdf, .docx, .txt"
        )

    try:
        docs = loader_fn(file_path)
    except Exception as exc:
        raise FileLoadError(f"Failed to read '{file_path.name}': {exc}") from exc

    if not docs:
        raise FileLoadError(
            f"'{file_path.name}' contains no extractable text "
            "(it may be empty or a scanned/image-only PDF)."
        )

    logger.info("Loaded %d section(s) from %s", len(docs), file_path.name)
    return docs
