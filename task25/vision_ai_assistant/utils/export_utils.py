"""
utils/export_utils.py
=====================
Export pipeline: convert extraction results to JSON, Markdown, PDF, DOCX, TXT.

Each function accepts an ExtractionResult + ChatSession and returns
(bytes, filename, mime_type) ready for st.download_button().
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from loguru import logger

from models.document import ExtractionResult
from models.chat import ChatSession


# ---------------------------------------------------------------------------
# Type alias
# ---------------------------------------------------------------------------
ExportOutput = Tuple[bytes, str, str]   # (data, filename, mime_type)


# ---------------------------------------------------------------------------
# JSON export
# ---------------------------------------------------------------------------

def export_json(
    result: ExtractionResult,
    session: Optional[ChatSession] = None,
) -> ExportOutput:
    """
    Export structured extraction as a formatted JSON file.
    """
    payload = result.to_export_dict()

    if session:
        payload["conversation"] = [
            {
                "role": m.role,
                "content": m.content,
                "timestamp": m.timestamp.isoformat(),
            }
            for m in session.messages
            if m.role != "system"
        ]

    json_bytes = json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")
    filename = _safe_filename(result.image.metadata.filename, "json")
    return json_bytes, filename, "application/json"


# ---------------------------------------------------------------------------
# Markdown export
# ---------------------------------------------------------------------------

def export_markdown(
    result: ExtractionResult,
    session: Optional[ChatSession] = None,
) -> ExportOutput:
    """
    Export a full Markdown report of the document analysis.
    """
    from config.constants import DOCUMENT_TYPE_LABELS, DOCUMENT_TYPE_ICONS

    doc_type  = result.document_type
    icon      = DOCUMENT_TYPE_ICONS.get(doc_type, "📄")
    type_label = DOCUMENT_TYPE_LABELS.get(doc_type, "Document")
    now        = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    lines = [
        f"# {icon} Vision AI Assistant — Analysis Report",
        f"",
        f"**Generated:** {now}  ",
        f"**Document:** `{result.image.metadata.filename}`  ",
        f"**Type:** {type_label}  ",
        f"**Language:** {result.analysis.language_detected}  ",
        f"**Model:** {result.analysis.model_used}  ",
        f"",
        "---",
        "",
        "## 📋 Document Summary",
        "",
        result.analysis.initial_summary or "_No summary available._",
        "",
        "---",
        "",
    ]

    # Structured extraction
    if result.raw_extraction:
        lines += [
            "## 🗂️ Extracted Information",
            "",
            "```json",
            json.dumps(result.raw_extraction, indent=2, ensure_ascii=False),
            "```",
            "",
            "---",
            "",
        ]

    # Conversation history
    if session and session.messages:
        lines += ["## 💬 Conversation History", ""]
        for msg in session.messages:
            if msg.role == "system":
                continue
            prefix = "**You:**" if msg.is_user else "**Assistant:**"
            lines.append(f"{prefix}  ")
            lines.append(msg.content)
            lines.append("")

    md_text  = "\n".join(lines)
    md_bytes = md_text.encode("utf-8")
    filename = _safe_filename(result.image.metadata.filename, "md")
    return md_bytes, filename, "text/markdown"


# ---------------------------------------------------------------------------
# TXT export
# ---------------------------------------------------------------------------

def export_txt(
    result: ExtractionResult,
    session: Optional[ChatSession] = None,
) -> ExportOutput:
    """
    Export a plain-text version of the analysis.
    """
    now  = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    lines = [
        "VISION AI ASSISTANT — ANALYSIS REPORT",
        "=" * 50,
        f"Generated   : {now}",
        f"Document    : {result.image.metadata.filename}",
        f"Type        : {result.document_type}",
        f"Language    : {result.analysis.language_detected}",
        "=" * 50,
        "",
        "SUMMARY",
        "-" * 30,
        result.analysis.initial_summary or "No summary available.",
        "",
    ]

    if result.raw_extraction:
        lines += [
            "EXTRACTED DATA",
            "-" * 30,
            json.dumps(result.raw_extraction, indent=2, ensure_ascii=False),
            "",
        ]

    if session and session.messages:
        lines += ["CONVERSATION", "-" * 30]
        for msg in session.messages:
            if msg.role == "system":
                continue
            role_label = "You" if msg.is_user else "Assistant"
            lines.append(f"[{role_label}]")
            lines.append(msg.content)
            lines.append("")

    txt_bytes = "\n".join(lines).encode("utf-8")
    filename  = _safe_filename(result.image.metadata.filename, "txt")
    return txt_bytes, filename, "text/plain"


# ---------------------------------------------------------------------------
# PDF export (using reportlab)
# ---------------------------------------------------------------------------

def export_pdf(
    result: ExtractionResult,
    session: Optional[ChatSession] = None,
) -> ExportOutput:
    """
    Export a formatted PDF report using ReportLab.
    Falls back to a plain-text PDF if ReportLab is unavailable.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            HRFlowable, Preformatted, Table, TableStyle,
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "Title2",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=colors.HexColor("#6366f1"),
            spaceAfter=6,
        )
        h2_style = ParagraphStyle(
            "H2",
            parent=styles["Heading2"],
            fontSize=13,
            textColor=colors.HexColor("#4f46e5"),
            spaceBefore=12,
            spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "Body2",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=4,
        )
        meta_style = ParagraphStyle(
            "Meta",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#64748b"),
        )
        code_style = ParagraphStyle(
            "Code",
            parent=styles["Code"],
            fontSize=8,
            leading=11,
            backColor=colors.HexColor("#f1f5f9"),
            leftIndent=10,
        )

        now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        story = []

        # Title
        story.append(Paragraph("Vision AI Assistant — Analysis Report", title_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#6366f1")))
        story.append(Spacer(1, 0.3*cm))

        # Meta block
        meta_lines = [
            f"<b>Generated:</b> {now}",
            f"<b>Document:</b> {result.image.metadata.filename}",
            f"<b>Type:</b> {result.document_type}",
            f"<b>Language:</b> {result.analysis.language_detected}",
            f"<b>Model:</b> {result.analysis.model_used}",
        ]
        for line in meta_lines:
            story.append(Paragraph(line, meta_style))
        story.append(Spacer(1, 0.4*cm))

        # Summary
        story.append(Paragraph("Document Summary", h2_style))
        summary_text = result.analysis.initial_summary or "No summary available."
        # Escape HTML entities for ReportLab
        safe = summary_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, body_style))
        story.append(Spacer(1, 0.3*cm))

        # Extracted data
        if result.raw_extraction:
            story.append(Paragraph("Extracted Information", h2_style))
            json_str = json.dumps(result.raw_extraction, indent=2, ensure_ascii=False)
            # Limit to 80 chars per line to fit page
            story.append(Preformatted(json_str, code_style))
            story.append(Spacer(1, 0.3*cm))

        # Conversation
        if session and session.messages:
            story.append(Paragraph("Conversation History", h2_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1")))
            story.append(Spacer(1, 0.2*cm))

            for msg in session.messages:
                if msg.role == "system":
                    continue
                role_label = "You" if msg.is_user else "Assistant"
                role_color = "#1e3a5f" if msg.is_user else "#1a1a2e"
                label_style = ParagraphStyle(
                    f"label_{msg.id}",
                    parent=styles["Normal"],
                    fontSize=9,
                    textColor=colors.HexColor("#6366f1"),
                    fontName="Helvetica-Bold",
                )
                story.append(Paragraph(role_label, label_style))
                safe_content = msg.content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_content, body_style))
                story.append(Spacer(1, 0.15*cm))

        doc.build(story)
        pdf_bytes = buf.getvalue()
        filename  = _safe_filename(result.image.metadata.filename, "pdf")
        return pdf_bytes, filename, "application/pdf"

    except ImportError:
        logger.warning("ReportLab not available — falling back to TXT export")
        txt_bytes, filename, _ = export_txt(result, session)
        return txt_bytes, filename.replace(".txt", ".pdf"), "text/plain"
    except Exception as exc:
        logger.error("PDF export failed: {}", exc)
        txt_bytes, filename, _ = export_txt(result, session)
        return txt_bytes, filename.replace(".txt", "_fallback.txt"), "text/plain"


# ---------------------------------------------------------------------------
# DOCX export (using python-docx)
# ---------------------------------------------------------------------------

def export_docx(
    result: ExtractionResult,
    session: Optional[ChatSession] = None,
) -> ExportOutput:
    """
    Export a formatted Word document using python-docx.
    Falls back to TXT if python-docx is unavailable.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

        # Title
        title = doc.add_heading("Vision AI Assistant — Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_heading_color(title, (99, 102, 241))  # indigo

        doc.add_paragraph()

        # Meta table
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = "Table Grid"
        meta_data = [
            ("Generated", now),
            ("Document", result.image.metadata.filename),
            ("Type", result.document_type),
            ("Language", result.analysis.language_detected),
            ("Model", result.analysis.model_used),
        ]
        for i, (key, val) in enumerate(meta_data):
            meta_table.cell(i, 0).text = key
            meta_table.cell(i, 1).text = val or ""

        doc.add_paragraph()

        # Summary
        doc.add_heading("Document Summary", level=1)
        doc.add_paragraph(result.analysis.initial_summary or "No summary available.")

        # Extracted data
        if result.raw_extraction:
            doc.add_heading("Extracted Information", level=1)
            code_para = doc.add_paragraph(
                json.dumps(result.raw_extraction, indent=2, ensure_ascii=False)
            )
            code_para.style = "No Spacing"
            run = code_para.runs[0] if code_para.runs else code_para.add_run("")
            run.font.name = "Courier New"
            run.font.size = Pt(8)

        # Conversation
        if session and session.messages:
            doc.add_heading("Conversation History", level=1)
            for msg in session.messages:
                if msg.role == "system":
                    continue
                role_label = "You" if msg.is_user else "Assistant"
                para = doc.add_paragraph()
                run  = para.add_run(f"{role_label}: ")
                run.bold = True
                run.font.color.rgb = RGBColor(99, 102, 241)
                para.add_run(msg.content)

        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        filename   = _safe_filename(result.image.metadata.filename, "docx")
        return docx_bytes, filename, (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        )

    except ImportError:
        logger.warning("python-docx not available — falling back to TXT")
        txt_bytes, filename, _ = export_txt(result, session)
        return txt_bytes, filename.replace(".txt", ".docx"), "text/plain"
    except Exception as exc:
        logger.error("DOCX export failed: {}", exc)
        txt_bytes, filename, _ = export_txt(result, session)
        return txt_bytes, filename.replace(".txt", "_fallback.txt"), "text/plain"


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def export(
    fmt: str,
    result: ExtractionResult,
    session: Optional[ChatSession] = None,
) -> ExportOutput:
    """
    Unified export dispatcher.

    Args:
        fmt:     One of 'json', 'markdown', 'pdf', 'docx', 'txt'
        result:  ExtractionResult from the pipeline
        session: Optional ChatSession for conversation history

    Returns:
        (bytes, filename, mime_type)
    """
    dispatch = {
        "json":     export_json,
        "markdown": export_markdown,
        "pdf":      export_pdf,
        "docx":     export_docx,
        "txt":      export_txt,
    }
    fn = dispatch.get(fmt.lower(), export_txt)
    logger.info("Exporting as {} | {}", fmt, result.image.metadata.filename)
    return fn(result, session)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _safe_filename(original: str, ext: str) -> str:
    """Build a safe export filename from the original image filename."""
    stem = Path(original).stem
    # Remove unsafe characters
    safe = "".join(c if c.isalnum() or c in "-_." else "_" for c in stem)
    ts   = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    return f"{safe}_{ts}.{ext}"


def _set_heading_color(heading, rgb: tuple) -> None:
    """Set all runs in a docx heading to an RGB color."""
    try:
        from docx.shared import RGBColor
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*rgb)
    except Exception:
        pass
