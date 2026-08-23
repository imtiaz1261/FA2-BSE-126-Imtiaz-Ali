"""
Document processing services.

Handles file extraction, text chunking, and metadata extraction for various file types:
- PDF (PyPDF2)
- DOCX (python-docx)
- CSV (csv module)
- TXT (plain text)
"""

import csv
import io
import logging
from abc import ABC, abstractmethod
from typing import Optional

import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# ============================================================================
# Configuration
# ============================================================================

SUPPORTED_FILE_TYPES = {"pdf", "docx", "txt", "csv"}
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20MB default

# Token estimation: ~4 characters per token on average
CHARS_PER_TOKEN = 4
TARGET_CHUNK_SIZE_TOKENS = 500
TARGET_CHUNK_SIZE_CHARS = TARGET_CHUNK_SIZE_TOKENS * CHARS_PER_TOKEN  # ~2000 chars
CHUNK_OVERLAP_TOKENS = 100
CHUNK_OVERLAP_CHARS = CHUNK_OVERLAP_TOKENS * CHARS_PER_TOKEN  # ~400 chars


# ============================================================================
# File Loader Interface
# ============================================================================


class FileLoader(ABC):
    """Base class for file loaders."""

    @abstractmethod
    def extract_text(self, file_content: bytes) -> tuple[str, dict]:
        """
        Extract text from file.

        Args:
            file_content: Raw file bytes

        Returns:
            (text, metadata) where metadata may contain page numbers, etc.
        """
        pass

    def validate_file_size(self, file_size: int) -> None:
        """Validate file size."""
        if file_size > MAX_FILE_SIZE_BYTES:
            raise ValueError(
                f"File size {file_size} exceeds maximum of {MAX_FILE_SIZE_BYTES} bytes"
            )


# ============================================================================
# PDF Loader
# ============================================================================


class PDFLoader(FileLoader):
    """Load and extract text from PDF files."""

    def extract_text(self, file_content: bytes) -> tuple[str, dict]:
        """Extract text and page numbers from PDF."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            page_numbers = []

            for page_num, page in enumerate(pdf_reader.pages, start=1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        # Track page numbers for later metadata
                        text_parts.append(text)
                        page_numbers.append(page_num)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num}: {e}")
                    continue

            full_text = "\n\n".join(text_parts)
            metadata = {
                "total_pages": len(pdf_reader.pages),
                "page_numbers": page_numbers,
            }

            if not full_text.strip():
                raise ValueError("No extractable text found in PDF")

            return full_text, metadata

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")


# ============================================================================
# DOCX Loader
# ============================================================================


class DocxLoader(FileLoader):
    """Load and extract text from DOCX files."""

    def extract_text(self, file_content: bytes) -> tuple[str, dict]:
        """Extract text from DOCX."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)
            metadata = {"total_paragraphs": len(paragraphs), "has_tables": len(doc.tables) > 0}

            if not full_text.strip():
                raise ValueError("No extractable text found in DOCX")

            return full_text, metadata

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")


# ============================================================================
# CSV Loader
# ============================================================================


class CSVLoader(FileLoader):
    """Load and extract text from CSV files."""

    def extract_text(self, file_content: bytes) -> tuple[str, dict]:
        """Extract text from CSV, treating as structured data."""
        try:
            # Decode with UTF-8, fallback to latin-1
            try:
                text_content = file_content.decode("utf-8")
            except UnicodeDecodeError:
                text_content = file_content.decode("latin-1")

            reader = csv.DictReader(io.StringIO(text_content))
            rows = list(reader)

            if not rows:
                raise ValueError("CSV is empty")

            # Build readable text from rows
            parts = []
            headers = reader.fieldnames or []

            for i, row in enumerate(rows, start=1):
                row_text = f"Row {i}: " + " | ".join(
                    f"{k}: {v}" for k, v in row.items() if v
                )
                parts.append(row_text)

            full_text = "\n".join(parts)
            metadata = {
                "total_rows": len(rows),
                "headers": headers,
                "total_columns": len(headers),
            }

            return full_text, metadata

        except Exception as e:
            logger.error(f"CSV extraction error: {e}")
            raise ValueError(f"Failed to process CSV: {str(e)}")


# ============================================================================
# Plain Text Loader
# ============================================================================


class TextLoader(FileLoader):
    """Load plain text files."""

    def extract_text(self, file_content: bytes) -> tuple[str, dict]:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode("utf-8")
            except UnicodeDecodeError:
                text = file_content.decode("latin-1")

            if not text.strip():
                raise ValueError("Text file is empty")

            # Count lines
            lines = text.split("\n")
            metadata = {"total_lines": len(lines)}

            return text, metadata

        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            raise ValueError(f"Failed to process text file: {str(e)}")


# ============================================================================
# Loader Factory
# ============================================================================


def get_loader(file_type: str) -> FileLoader:
    """Get appropriate loader for file type."""
    file_type = file_type.lower().strip(".")

    if file_type == "pdf":
        return PDFLoader()
    elif file_type == "docx":
        return DocxLoader()
    elif file_type == "csv":
        return CSVLoader()
    elif file_type == "txt":
        return TextLoader()
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def validate_file_type(file_type: str) -> bool:
    """Validate if file type is supported."""
    return file_type.lower().strip(".") in SUPPORTED_FILE_TYPES


# ============================================================================
# Text Chunking
# ============================================================================


def chunk_text(
    text: str,
    chunk_size_tokens: int = TARGET_CHUNK_SIZE_TOKENS,
    overlap_tokens: int = CHUNK_OVERLAP_TOKENS,
) -> list[tuple[str, int]]:
    """
    Chunk text into overlapping chunks of approximately chunk_size_tokens.

    Uses character-based estimation for token counts.

    Args:
        text: Full text to chunk
        chunk_size_tokens: Target chunk size in tokens (~500)
        overlap_tokens: Overlap between chunks (~100)

    Returns:
        List of (chunk_text, token_count) tuples
    """
    chunk_size_chars = chunk_size_tokens * CHARS_PER_TOKEN
    overlap_chars = overlap_tokens * CHARS_PER_TOKEN

    chunks = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size_chars, len(text))

        # Try to break at a sentence boundary for better chunking
        if end < len(text):
            # Look for sentence boundary (. followed by space) within last 200 chars
            search_start = max(start, end - 200)
            last_period = text.rfind(". ", search_start, end)
            if last_period != -1 and last_period > start:
                end = last_period + 2  # Include the period and space

        chunk = text[start:end].strip()

        if chunk:
            # Estimate token count
            token_count = len(chunk) // CHARS_PER_TOKEN
            chunks.append((chunk, token_count))

        # Move start with overlap
        start = end - overlap_chars

        # Prevent infinite loop with very small text
        if start >= end - overlap_chars:
            start = end

    return chunks


def estimate_token_count(text: str) -> int:
    """Estimate token count from character count."""
    return len(text) // CHARS_PER_TOKEN


# ============================================================================
# Full Document Processing
# ============================================================================


class DocumentProcessor:
    """High-level document processor combining extraction and chunking."""

    @staticmethod
    def process(
        file_content: bytes,
        filename: str,
        file_type: str,
        max_chunks: Optional[int] = None,
    ) -> tuple[list[tuple[str, int]], dict, Optional[str]]:
        """
        Process a document: extract text, validate, and chunk.

        Args:
            file_content: Raw file bytes
            filename: Original filename
            file_type: File extension (pdf, docx, txt, csv)
            max_chunks: Optional limit on number of chunks to return

        Returns:
            (chunks, metadata, error_message)
            - chunks: List of (text, token_count) tuples
            - metadata: Dict with extraction metadata (page numbers, etc)
            - error_message: Error string if processing failed, else None
        """
        try:
            # Validate file type
            if not validate_file_type(file_type):
                return [], {}, f"Unsupported file type: {file_type}"

            # Validate file size
            if len(file_content) > MAX_FILE_SIZE_BYTES:
                return [], {}, (
                    f"File size {len(file_content)} bytes exceeds "
                    f"maximum of {MAX_FILE_SIZE_BYTES} bytes"
                )

            # Extract text
            loader = get_loader(file_type)
            text, metadata = loader.extract_text(file_content)

            # Chunk text
            chunks = chunk_text(text)

            if max_chunks and len(chunks) > max_chunks:
                chunks = chunks[:max_chunks]

            logger.info(
                f"Processed {filename} ({file_type}): {len(chunks)} chunks, "
                f"{len(text)} chars"
            )

            return chunks, metadata, None

        except Exception as e:
            error_msg = f"Document processing error: {str(e)}"
            logger.error(error_msg)
            return [], {}, error_msg
